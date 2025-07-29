import traceback
import io

import base64
from typing import List

import docx
from docx.shared import Inches, Pt

import tempfile
import os

import markdown

try:
    from pydocmaker.backend.baseformatter import BaseFormatter
except Exception as err:
    from .baseformatter import BaseFormatter

can_run_pandoc = lambda : False


try:
    from pydocmaker.backend.pandoc_api import can_run_pandoc, pandoc_convert, pandoc_convert_file
except Exception as err:
    from .pandoc_api import can_run_pandoc, pandoc_convert, pandoc_convert_file

try:
    from pydocmaker.backend.ex_html import convert as convert_html
except Exception as err:
    from .ex_html import convert as convert_html



def blue(run):
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

def red(run):
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

def convert_pandoc(doc:List[dict]) -> bytes:

    with tempfile.TemporaryDirectory() as temp_dir:
        html_file_path = os.path.join(temp_dir, 'temp.html')
        docx_file_path = os.path.join(temp_dir, 'temp.docx')

        with open(html_file_path, 'w', encoding='utf-8') as fp:
            fp.write(convert_html(doc))
        
        pandoc_convert_file(html_file_path, docx_file_path)
        with open(docx_file_path, 'rb') as fp:
            return fp.read()
        
def convert(doc:List[dict]) -> bytes:

    if can_run_pandoc():
        return convert_pandoc(doc)
    else:
        renderer = docx_renderer()
        renderer.digest(doc)
        return renderer.doc_to_bytes()

class docx_renderer(BaseFormatter):
    def __init__(self, template_path:str=None, make_blue=False) -> None:
        self.d = docx.Document(template_path)
        self.make_blue = make_blue

    def add_paragraph(self, newtext, *args, **kwargs):
        new_paragraph = self.d.add_paragraph(newtext, *args, **kwargs)
        if self.make_blue:
            for r in new_paragraph.runs:
                blue(r)
        return new_paragraph
    
    def add_run(self, text, *args, **kwargs):
        if not self.d.paragraphs:
            self.add_paragraph('')

        last_paragraph = self.d.paragraphs[-1]
        
        if not last_paragraph.runs:
            last_run = last_paragraph.add_run(text)
        else:
            last_run = last_paragraph.runs[-1]
            last_run.add_text(text)
            
        if self.make_blue:
            blue(last_run)
        return last_run
        
    def digest_text(self, children, *args, **kwargs):
        return self.add_paragraph(children)
    

    def digest_str(self, children, *args, **kwargs):
        return self.add_run(children)

    def digest_line(self, children, *args, **kwargs):
        return self.add_run(children + '\n')
    
    def digest_markdown(self, children, *args, **kwargs):
        return self.add_paragraph(children, style='Normal')
        
    def digest_verbatim(self, children, *args, **kwargs):
        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        return new_run

    def digest_latex(self, children, *args, **kwargs):
        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        return new_run


    def handle_error(self, err, el=None) -> list:
        if isinstance(err, BaseException):
            traceback.print_exc(limit=5)
            err = '\n'.join(traceback.format_exception(type(err), value=err, tb=err.__traceback__, limit=5))

        new_run = self.add_run(err)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        red(new_run)
        return new_run


    def digest_iterator(self, children, *args, **kwargs):
        if children:
            return [self.digest(val, *args, **kwargs) for val in children]
        return []

    def digest_table(self, children=None, **kwargs) -> str:
        self.handle_error(NotImplementedError(f'exporter of type {type(self)} can not handle tables'))
    
    def digest_image(self, children, *args, **kwargs):

        image_width = Inches(max(1, kwargs.get('width', 0.8)*5))
        image_caption = kwargs.get('caption', '')
        image_blob = kwargs.get('imageblob', '')

        assert image_blob, 'no image data given!'

        btsb64 = image_blob.split(',')[-1]

        # Decode the base64 image
        img_bytes = base64.b64decode(btsb64)

        # Create an image stream from the bytes
        image_stream = io.BytesIO(img_bytes)
        
        picture = self.d.add_picture(image_stream, width=image_width)
        # picture.width = image_width  # Ensure fixed width
        # picture.height = None  # Adjust height automatically
        picture.alignment = 1

        run = self.add_paragraph(image_caption)
        # run.style = 'Caption'  # Apply the 'Caption' style for formatting

        return run

    def format(self, *args, **kwargs):
        raise NotImplementedError('Can not format a docx document directly')
    

    def doc_to_bytes(self):
        with io.BytesIO() as fp:
            self.d.save(fp)
            fp.seek(0)
            return fp.read()

    def save(self, filepath):
        self.d.save(filepath)
