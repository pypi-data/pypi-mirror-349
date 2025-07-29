import logging
from typing import List, Tuple

from docx import Document

from .document_chunking import BaseDocumentChunker, TOCNode

# Get logger for this module
logger = logging.getLogger(__name__)


class DOCXTOCChunker(BaseDocumentChunker):
    """
    A document chunker that creates a hierarchical tree of nodes based on Word document headings.
    """

    def __init__(self, docx_path: str, source_display_name: str):
        """
        Initialize the chunker with the path to the Word document.

        Args:
            docx_path: Path to the Word document file
            source_display_name: The original name of the source
        """
        super().__init__(docx_path, source_display_name)
        self.doc = None
        self.paragraphs = []
        self._document_loaded = False
        self._heading_style_to_level = {
            "Heading 1": 1,
            "Heading 2": 2,
            "Heading 3": 3,
            "Heading 4": 4,
            "Heading 5": 5,
            "Heading 6": 6,
            "Title": 0,
        }

    def load_document(self) -> None:
        """Load the Word document and extract its structure"""
        try:
            self.doc = Document(self.source_path)
            self.paragraphs = list(self.doc.paragraphs)
            self._document_loaded = True
        except Exception as e:
            logger.error(f"Error loading Word document: {e}")
            raise

    def build_toc_tree(self) -> TOCNode:
        """
        Build a tree structure from the Word document headings.

        Returns:
            The root node of the TOC tree
        """
        if not self._document_loaded:
            self.load_document()

        # Initialize the document tree
        self.root_node = TOCNode(title="Document Root", page_num=0, level=0)

        # Find all headings in the document
        headings = self._extract_headings()

        if not headings:
            # No headings found, treat the entire document as one chunk
            self.root_node.content = self._extract_all_text()
            self.root_node.end_page = 0  # DOCX doesn't have page numbers in Python API
            return self.root_node

        # Process headings and build TOC tree
        for i, (para_idx, level, title) in enumerate(headings):
            # Create a new node for this header
            node = TOCNode(
                title=title,
                page_num=para_idx,  # Using paragraph index instead of page number
                level=level,
            )

            # Find the appropriate parent for this node based on its level
            parent = self._find_parent_for_level(self.root_node, level)
            parent.add_child(node)

            # Determine the content range for this node
            start_idx = para_idx + 1  # Start after the header paragraph

            # End index is either the index before the next heading or the end of document
            end_idx = len(self.paragraphs)
            if i < len(headings) - 1:
                end_idx = headings[i + 1][0]  # Paragraph index of next header

            # Extract content for this section
            node.content = self._extract_content(start_idx, end_idx)
            node.end_page = 0  # DOCX doesn't have page numbers in Python API

        return self.root_node

    def _extract_headings(self) -> List[Tuple[int, int, str]]:
        """
        Extract heading paragraphs from the Word document.

        Returns:
            List of tuples with (paragraph_index, heading_level, heading_title)
        """
        headings = []

        for i, para in enumerate(self.paragraphs):
            if para.style.name in self._heading_style_to_level:
                level = self._heading_style_to_level[para.style.name]
                title = para.text.strip()
                if title:  # Only include headings with text content
                    headings.append((i, level, title))

        return headings

    def _find_parent_for_level(self, node: TOCNode, target_level: int) -> TOCNode:
        """
        Find the appropriate parent node for a node with the given level.

        Args:
            node: Current node to start search from
            target_level: The level of the node we want to find a parent for

        Returns:
            The appropriate parent node
        """
        # If this is a level 1 header, its parent is the root node
        if target_level == 1:
            return self.root_node

        # If current node is at the level just above our target, it's our parent
        if node.level == target_level - 1:
            return node

        # If node has children, check the last child first (most recent node at that level)
        if node.children:
            last_child = node.children[-1]
            # If the last child's level is less than our target level, it could be our parent
            if last_child.level < target_level:
                return self._find_parent_for_level(last_child, target_level)

        # If we're here, we need to move up the tree
        if node.parent:
            return self._find_parent_for_level(node.parent, target_level)

        # Fallback to root node if no appropriate parent is found
        return self.root_node

    def _extract_content(self, start_idx: int, end_idx: int) -> str:
        """
        Extract content from the specified paragraph range.

        Args:
            start_idx: Starting paragraph index (inclusive)
            end_idx: Ending paragraph index (exclusive)

        Returns:
            The extracted content as a string
        """
        if start_idx >= end_idx or start_idx >= len(self.paragraphs):
            return ""

        # Adjust end_idx if it exceeds the document length
        end_idx = min(end_idx, len(self.paragraphs))

        # Extract paragraphs in the range
        content_paragraphs = []
        for i in range(start_idx, end_idx):
            para = self.paragraphs[i]
            if para.text.strip():  # Only include non-empty paragraphs
                content_paragraphs.append(para.text)

        return "\n".join(content_paragraphs)

    def _extract_all_text(self) -> str:
        """Extract all text content from the document"""
        return "\n".join([p.text for p in self.paragraphs if p.text.strip()])

    def close(self) -> None:
        """Clean up resources"""
        self.doc = None
        self.paragraphs = []
        self._document_loaded = False
