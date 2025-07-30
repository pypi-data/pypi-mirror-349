"""
DOCX file handler module for reading and writing Microsoft Word documents.
"""

import os
from cleanfusion.utils.logger import Logger

class DOCXHandler:
    """
    Handler for Microsoft Word (.docx) files.
    """
    
    def __init__(self):
        """Initialize the DOCX handler."""
        self.logger = Logger()
    
    def read_file(self, file_path):
        """
        Extract text from a DOCX file.
        
        Parameters
        ----------
        file_path : str
            Path to the DOCX file.
        
        Returns
        -------
        str
            The extracted text from the DOCX.
        """
        try:
            self.logger.info(f"Reading DOCX file: {file_path}")
            
            try:
                import docx
                doc = docx.Document(file_path)
                
                # Extract text from paragraphs
                paragraphs = [p.text for p in doc.paragraphs]
                
                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        tables_text.append(' | '.join(row_text))
                
                # Combine all text
                text = '\n\n'.join(paragraphs)
                if tables_text:
                    text += '\n\n' + '\n'.join(tables_text)
                
                self.logger.info(f"Successfully extracted text from DOCX: {len(paragraphs)} paragraphs")
                return text
                
            except ImportError:
                self.logger.warning("python-docx not installed. Attempting to use docx2txt...")
                try:
                    import docx2txt
                    text = docx2txt.process(file_path)
                    self.logger.info(f"Successfully extracted text from DOCX using docx2txt")
                    return text
                except ImportError:
                    self.logger.error("Neither python-docx nor docx2txt is installed. Please install one of them.")
                    raise ImportError("DOCX extraction requires python-docx or docx2txt. Please install one of them.")
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {e}")
            raise
    
    def write_file(self, text, file_path):
        """
        Write text to a DOCX file.
        
        Parameters
        ----------
        text : str
            The text to write.
        
        file_path : str
            Path to the output DOCX file.
        
        Returns
        -------
        str
            The path to the written file.
        """
        try:
            self.logger.info(f"Writing DOCX file: {file_path}")
            
            try:
                import docx
                doc = docx.Document()
                
                # Split text into paragraphs
                paragraphs = text.split('\n')
                
                # Add each paragraph to the document
                for paragraph in paragraphs:
                    if paragraph.strip():  # Skip empty paragraphs
                        doc.add_paragraph(paragraph)
                
                # Save the document
                doc.save(file_path)
                
                self.logger.info(f"Successfully wrote DOCX file with {len(paragraphs)} paragraphs")
                return file_path
                
            except ImportError:
                self.logger.error("python-docx not installed. Cannot write DOCX file.")
                raise ImportError("Writing DOCX files requires python-docx. Please install it.")
            
        except Exception as e:
            self.logger.error(f"Error writing DOCX file: {e}")
            raise
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from a DOCX file.
        
        Parameters
        ----------
        file_path : str
            Path to the DOCX file.
        
        Returns
        -------
        dict
            Dictionary containing document metadata.
        """
        try:
            self.logger.info(f"Extracting metadata from DOCX file: {file_path}")
            
            try:
                import docx
                doc = docx.Document(file_path)
                
                # Extract core properties
                core_props = doc.core_properties
                
                metadata = {
                    'author': core_props.author,
                    'created': core_props.created,
                    'last_modified_by': core_props.last_modified_by,
                    'modified': core_props.modified,
                    'title': core_props.title,
                    'subject': core_props.subject,
                    'keywords': core_props.keywords,
                    'category': core_props.category,
                    'comments': core_props.comments,
                    'content_status': core_props.content_status,
                    'language': core_props.language,
                    'version': core_props.revision
                }
                
                self.logger.info(f"Successfully extracted metadata from DOCX file")
                return metadata
                
            except ImportError:
                self.logger.error("python-docx not installed. Cannot extract metadata.")
                raise ImportError("Extracting metadata from DOCX files requires python-docx. Please install it.")
            
        except Exception as e:
            self.logger.error(f"Error extracting metadata from DOCX file: {e}")
            raise