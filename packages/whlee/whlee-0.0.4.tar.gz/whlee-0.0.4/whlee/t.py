import datetime, time

def timee(func):  # decorator for calculate the wall time of a function and print.
    def w1(*args, **kw):
        start_time = datetime.datetime.now()
        rst = func(*args, **kw)
        end_time = datetime.datetime.now()
        interval = (end_time - start_time).seconds
        print(time.strftime("%H:%M:%S", time.gmtime(interval)) )
        return rst
    return w1

import io, collections, json, docx, pandas as pd, numpy as np
# ________________________________________________________________________


class TitleMark:
    """
    dc = docx.Document()
    tm = TitleMark(dc)
    tm.addhead1('示例')  # 使用默认编号，不同级别自动选择
    tm.addhead1('示例', mark_type='empty')  # 标题前面不使用编号
    tm.addhead1('示例', mark_type='default_level_5')  # 指定特定的编号，可随意选择现有的编号格式
    tm.add_pic_note('示例')  # 图片说明，按章从1计数。放在图片底部，之后有一空行。
    tm.add_table_note('示例')  # 表格说明，按章从1计数。放在表格顶部，之前有一空行。
    # 支持另一种接口方式
    tm.addhead('示例', level=1)

    example
    ------------------------------------------------
    import docx

    dc = docx.Document()
    style1 = {'font_name': '方正小标宋_GBK', 'font_size' : 24, 'pf_alignment' : 1}
    TitleMark.addstyle(dc, 's1', **style1)  # 封面主标题
    style1 = {'font_name': '方正黑体_GBK', 'font_size' : 14, 'pf_alignment' : 1}
    TitleMark.addstyle(dc, 's1_1', **style1)  # 封面副标题
    style1 = {'font_name': '方正楷体_GBK', 'font_size' : 10, 'pf_alignment' : 1}
    TitleMark.addstyle(dc, 's2', **style1)  # 图例说明
    style1 = {'font_name': '方正仿宋_GBK', 'font_size' : 14, 'pf_first_line_indent' : 30, 'pf_line_spacing' : 30, 'pf_space_before' : 0, 'pf_space_after' : 5, }
    TitleMark.addstyle(dc, 's3', **style1)  # 正文
    tm = TitleMark(dc)

    _ = [dc.add_paragraph(style='s1') for i in range(3)]
    dc.add_paragraph('按价位段自选投放策略自动化生成', style='s1')
    _ = [dc.add_paragraph(style='s1') for i in range(7)]
    dc.add_paragraph('执行报告', style='s1_1')
    dc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    text = '基本信息'
    tm.addhead_deprecated(text, level=1)
    text = '执行耗时'
    tm.addhead_deprecated(text, level=2)
    text = '历史投放数据'
    tm.addhead_deprecated(text, level=3)
    text = '共计耗时秒'
    dc.add_paragraph(style='s3', text=text)

    dc.add_picture('示意图.png', width=docx.shared.Cm(15))

    TitleMark.addtable(dc, iris.head())

    frame_corr = iris.corr()
    fig = px.imshow(img=frame_corr, color_continuous_scale='Viridis').update_layout(width=400, height=300, margin=dict(l=1, r=1, b=1, t=1), font=dict(size=8))
    TitleMark.addpic(dc, fig)

    fig = px.scatter(x=iris['a0'], y=iris['a1'], ).update_layout(width=200, height=120, margin=dict(l=1, r=1, b=1, t=150), xaxis_range=[0,10], yaxis_range=[0,10], font_size=8)
    TitleMark.addplotly(dc, fig)

    tm.add_pic_note('水箱工艺流程冷1')
    tm.add_table_note('水箱工艺流程2')
    tm.add_pic_note('水箱工艺流程冷3')
    tm.add_table_note('水箱工艺流程4')

    dc.save('test.docx')

    """
    def __init__(self, dc):
        self.dc = dc
        self.counter = collections.OrderedDict({1:0, 2:0, 3:0, 4:0, 5:0})
        self.pic_note_counter = 0
        self.table_note_counter = 0
        
        style_name = 'note'
        self.dc.styles.add_style(style_name, style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        self.dc.styles[style_name].font.name = ''
        self.dc.styles[style_name]._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '方正楷体_GBK')
        self.dc.styles[style_name].font.size = docx.shared.Pt(10)
        self.dc.styles[style_name].paragraph_format.alignment = 1
        
        self.marks_store = {}
        self.marks_store['default_level_1'] = {ind:f'{letter}、' for ind, letter in enumerate(list('一二三四五六七八九十')+['十一','十二','十三','十四','十五','十六'], start=1)}
        self.marks_store['default_level_2'] = {ind:f'（{letter}）' for ind, letter in enumerate(list('一二三四五六七八九十')+['十一','十二','十三','十四','十五','十六'], start=1)}
        self.marks_store['default_level_3'] = {i:f'{i}. ' for i in range(1, 100)}
        self.marks_store['default_level_4'] = {i:f'（{i}）' for i in range(1, 100)}
        self.marks_store['default_level_5'] = {ind:f'{letter}）' for ind, letter in enumerate('abcdefghijklmnopqrstuvwxyz', start=1)}
        self.marks_store['empty'] = {i:'' for i in range(1, 100)}

    def add_pic_note(self, text):
        self.pic_note_counter += 1
        text = f'图{self.counter[1]}-{self.pic_note_counter}.  {text}'
        self.dc.add_paragraph(text, style='note')
        self.dc.add_paragraph()
    def add_table_note(self, text):
        self.table_note_counter += 1
        text = f'表{self.counter[1]}-{self.table_note_counter}.  {text}'
        self.dc.add_paragraph()
        self.dc.add_paragraph(text, style='note')
    def addhead1(self, text, mark_type='default'):
        self.pic_note_counter = 0
        self.table_note_counter = 0
        self.counter[1] += 1
        self.counter[2] = 0
        self.counter[3] = 0
        self.counter[4] = 0
        self.counter[5] = 0
        self.add(level=1, text=text, font='方正黑体_GBK', size=16, bold=True, line_spacing=0, space_before=15, space_after=5, first_line_indent=0, mark_type=mark_type)
    def addhead2(self, text, mark_type='default'):
        self.counter[2] += 1
        self.counter[3] = 0
        self.counter[4] = 0
        self.counter[5] = 0
        self.add(level=2, text=text, font='方正楷体_GBK', size=16, bold=True, line_spacing=0, space_before=15, space_after=5, first_line_indent=0, mark_type=mark_type)
    def addhead3(self, text, mark_type='default'):
        self.counter[3] += 1
        self.counter[4] = 0
        self.counter[5] = 0
        self.add(level=3, text=text, font='方正仿宋_GBK', size=15, bold=True, line_spacing=0, space_before=15, space_after=5, first_line_indent=0, mark_type=mark_type)
    def addhead4(self, text, mark_type='default'):
        self.counter[4] += 1
        self.counter[5] = 0
        self.add(level=4, text=text, font='方正仿宋_GBK', size=14, bold=True, line_spacing=0, space_before=15, space_after=5, first_line_indent=0, mark_type=mark_type)
    def addhead5(self, text, mark_type='default'):
        self.counter[5] += 1
        self.add(level=5, text=text, font='方正仿宋_GBK', size=14, bold=True, line_spacing=0, space_before=15, space_after=5, first_line_indent=0, mark_type=mark_type)
    def add(self, **kwargs):
        parag = self.dc.add_heading(level=kwargs['level'])
        if kwargs['mark_type'] == 'default':
            if kwargs['level'] == 1:
                marks = self.marks_store['default_level_1']
            elif kwargs['level'] == 2:
                marks = self.marks_store['default_level_2']
            elif kwargs['level'] == 3:
                marks = self.marks_store['default_level_3']
            elif kwargs['level'] == 4:
                marks = self.marks_store['default_level_4']
            elif kwargs['level'] == 5:
                marks = self.marks_store['default_level_5']
        else:
            assert kwargs['mark_type'] in self.marks_store.keys(), f"the title mark {kwargs['mark_type']} is not supported"
            marks = self.marks_store[kwargs['mark_type']]
        counters_values = list(self.counter.values())
        curr_value = counters_values[kwargs['level']-1]
        mymark = marks.get(curr_value)
        text = mymark + kwargs['text']
        run = parag.add_run(text)
        run.font.name = ''
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), kwargs['font'])
        run.font.size = docx.shared.Pt(kwargs['size'])
        run.bold = kwargs['bold']
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        run.italic = False
        parag.paragraph_format.line_spacing_rule = 4
        parag.paragraph_format.line_spacing = docx.shared.Pt(kwargs['line_spacing'])
        parag.paragraph_format.space_before = docx.shared.Pt(kwargs['space_before'])
        parag.paragraph_format.space_after = docx.shared.Pt(kwargs['space_after'])
        parag.paragraph_format.first_line_indent  = docx.shared.Pt(kwargs['first_line_indent'])
    def addhead_deprecated(self, text, level):
        """兼容接口：旧的版本接口"""
        if level == 1:
            self.addhead1(text)
        elif level == 2:
            self.addhead2(text)
        elif level == 3:
            self.addhead3(text)
        elif level == 4:
            self.addhead4(text)
        elif level == 5:
            self.addhead5(text)

    @staticmethod
    def addpic(dc, fig):
        buf = io.BytesIO()
        fig.write_image(buf)
        p = dc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(buf)

    @staticmethod
    def addplotly(dc, fig, w=14, h=7):  # 根据图片的宽，自动设置图片的像素密度。只对plotly图片有效（update_layout）
        buf = io.BytesIO()
        fig.update_layout(font_size=w*1.8, width=w*100, height=h*100, margin=dict(l=2, r=2, b=2, t=2))
        fig.write_image(buf)
        p = dc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(buf, width=docx.shared.Cm(w), height=docx.shared.Cm(h))

    @staticmethod
    def addhead(dc, text, level, font='黑体', bold=True, space_line=50, space_before=1, space_after=1):
        parag = dc.add_heading(level=level)
        parag.paragraph_format.first_line_indent  = docx.shared.Pt(-5)
        parag.paragraph_format.line_spacing_rule = 4
        parag.paragraph_format.line_spacing = docx.shared.Pt(space_line)
        parag.paragraph_format.space_before = docx.shared.Pt(space_before)
        parag.paragraph_format.space_after = docx.shared.Pt(space_after)
        run = parag.add_run(text)
        run.font.name = ''
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font)
        if level == 1:
            run.font.size = docx.shared.Pt(24)
        elif level == 2:
            run.font.size = docx.shared.Pt(20)
        elif level == 3:
            run.font.size = docx.shared.Pt(16)
        elif level == 4:
            run.font.size = docx.shared.Pt(12)
        elif level == 5:
            run.font.size = docx.shared.Pt(10)
        run.bold = bold
        run.italic = False
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    @staticmethod
    def addstyle(dc, name, **kwargs):
        """
        {
        'font_name': '方正小标宋_GBC', 
         'font_size': 14, 
         'font_bold': False, 
         'font_italic': False, 
         'font_color_rgb': (0,222,0), 
         'pf_alignment': 1, 
         'pf_first_line_indent': 30, 
         'pf_line_spacing': 30, 
         'pf_space_before': 30, 
         'pf_space_after': 30, 
        }
        dc = docx.Document()
        style1 = {'font_size': 24}
        addstyle(dc, 's1', **style1)
        dc.add_paragraph(txt, style='s1')
        """
        dc.styles.add_style(name, style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        if 'font_name' in kwargs:
            dc.styles[name].font.name = ''
            dc.styles[name]._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), kwargs['font_name'])
        if 'font_size' in kwargs:
            dc.styles[name].font.size = docx.shared.Pt(kwargs['font_size'])
        if 'font_bold' in kwargs:
            dc.styles[name].font.bold = kwargs['font_bold']
        if 'font_italic' in kwargs:
            dc.styles[name].font.italic = kwargs['font_italic']
        if 'font_color_rgb' in kwargs:
            dc.styles[name].font.color.rgb = docx.shared.RGBColor(*kwargs['font_color_rgb'])
        if 'pf_alignment' in kwargs:
            dc.styles[name].paragraph_format.alignment = kwargs['pf_alignment']
        if 'pf_first_line_indent' in kwargs:
            dc.styles[name].paragraph_format.first_line_indent = docx.shared.Pt(kwargs['pf_first_line_indent'])
        if 'pf_line_spacing' in kwargs:
            dc.styles[name].paragraph_format.line_spacing_rule = 4
            dc.styles[name].paragraph_format.line_spacing = docx.shared.Pt(kwargs['pf_line_spacing'])
        if 'pf_space_before' in kwargs:
            dc.styles[name].paragraph_format.space_before = docx.shared.Pt(kwargs['pf_space_before'])
        if 'pf_space_after' in kwargs:
            dc.styles[name].paragraph_format.space_after = docx.shared.Pt(kwargs['pf_space_after'])

    # 添加表格，可设置字号、字体、列宽、首行背景、正文对齐
    @staticmethod
    def addtable(dc, df, size=11, font='方正仿宋_GBK', cwidth=None, headcolor='07CEFA', alignment=None,):
        shape = df.shape
        n_row = shape[0]
        n_col = shape[1]
        mytable = dc.add_table(n_row + 1, n_col)
        heading_cells = mytable.rows[0].cells
        for ind, col in enumerate(df):
            mycell = heading_cells[ind]
            mycell.text = col
            mycell.paragraphs[0].paragraph_format.alignment = 1
            mycell.vertical_alignment = 1
            run = mycell.paragraphs[0].runs[0]
            run.font.name = ''
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font)
            run.font.size = docx.shared.Pt(size)
            run.font.bold = True
            ele = docx.oxml.OxmlElement('w:shd')  # 创建阴影填充元素
            ele.set(docx.oxml.ns.qn('w:fill'), headcolor)  # 设置元素属性
            mycell._tc.get_or_add_tcPr().append(ele)
        for rowind in range(n_row):
            for colind in range(n_col):
                mycell = mytable.cell(rowind+1, colind)
                mycell.text = str(df.iat[rowind, colind])
                if alignment:
                    mycell.paragraphs[0].paragraph_format.alignment = alignment[colind]
                else:
                    mycell.paragraphs[0].paragraph_format.alignment = 1
                mycell.vertical_alignment = 1
                run = mycell.paragraphs[0].runs[0]
                run.font.name = ''
                run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font)
                run.font.size = docx.shared.Pt(size)
        if cwidth:
            for colind in range(n_col):
                col = mytable.columns[colind]
                col_width = cwidth[colind]
                for cell in col.cells:
                    cell.width = docx.shared.Cm(col_width)
        mytable.style = 'Table Grid'
        
