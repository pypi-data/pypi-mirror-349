from abc import ABC, abstractmethod
from pathlib import Path
from typing import Iterable, Literal, Union

from accel.base.box import Box
from accel.base.systems import System, Systems
from accel.util.log import logger

from acceltools.base import ToolBox


class DocOptions:
    def __init__(self) -> None:
        self.newline = None


class LayerAbc(ABC):
    def __init__(self) -> None:
        self.bold = False
        self.italic = False
        self.underline = False
        self.superscript = False
        self.subscript = False
        self.replace_hyphen = False
        super().__init__()

    def boldify(self):
        self.bold = True
        return self

    def italicize(self):
        self.italic = True
        return self

    @abstractmethod
    def get_str(self, c: System) -> str:
        pass


class Layers:
    class Name(LayerAbc):
        def __init__(self) -> None:
            super().__init__()

        def get_str(self, c: System) -> str:
            return c.name

    class Text(LayerAbc):
        def __init__(self, txt: str) -> None:
            super().__init__()
            self.txt = str(txt)

        def get_str(self, c: System) -> str:
            return self.txt

    class Data(LayerAbc):
        def __init__(self, key: str) -> None:
            super().__init__()
            self.key = key
            self.replace_hyphen = True

        def get_str(self, c: System) -> str:
            data = c.data.get(self.key)
            if data is None:
                logger.error(f"could not get {self.key} in {c.name}")
                data = "Nan"
            try:
                data = str(data)
            except TypeError:
                logger.error(f"could not convert the data {self.key} in {c.name}")
                data = "Nan"
            return data

    class LF(LayerAbc):
        def __init__(self, repeat: int = 1, linefeed="\n") -> None:
            super().__init__()
            self.repeat = repeat
            self.linefeed = linefeed

        def get_str(self, c: System) -> str:
            return "".join([self.linefeed for _ in range(self.repeat)])

    class Break(LayerAbc):
        def __init__(self, break_type: Literal["line", "page", "column", "section"] = "page") -> None:
            self.break_type = break_type
            super().__init__()

        def get_str(self, c: System) -> str:
            return ""

    class Energy(LayerAbc):
        def __init__(self, digit: int = 1, with_unit: bool = True) -> None:
            super().__init__()
            self.digit = int(digit)
            self.with_unit = bool(with_unit)
            self.replace_hyphen = True

        def get_str(self, c: System) -> str:
            fmt = f".{self.digit}f"
            if self.with_unit:
                rstr = f"{c.energy:{fmt}} kcal/mol"
            else:
                rstr = f"{c.energy:{fmt}}"
            return rstr

    class Cartesian(LayerAbc):
        def __init__(self, spacer=None, digit: int = 5, with_symbol: bool = True) -> None:
            super().__init__()
            self.digit = int(digit)
            self.with_symbol = bool(with_symbol)
            self.spacer = spacer
            self.replace_hyphen = True

        def get_str(self, c: System) -> str:
            rstr = ""
            if len(c.atoms) == 0:
                logger.error(f"no atoms in {c.name}")
                return rstr
            for a in c.atoms:
                if self.with_symbol:
                    if self.spacer is not None:
                        rstr += f"{a.symbol}{self.spacer}"
                    else:
                        rstr += f"{a.symbol:<3}"
                if self.spacer is not None:
                    rstr += f"{a.x:.{self.digit}f}{self.spacer}{a.y:.{self.digit}f}{self.spacer}{a.z:.{self.digit}f}\n"
                else:
                    fmt = f">{self.digit+5}.{self.digit}f"
                    rstr += f"{a.x:{fmt}}{a.y:{fmt}}{a.z:{fmt}}\n"
            return rstr


class DocBox(ToolBox):
    def __init__(self, contents: Union[Box, Systems, Iterable[System], System]):
        self.options = DocOptions()
        self.layers: list[LayerAbc] = []
        super().__init__(contents)

    def export_txt(self, path: Path, replace_hyphen: bool = False):
        rstr = ""
        for c in self.get():
            for lyr in self.layers:
                lyr_text = lyr.get_str(c)
                if replace_hyphen and lyr.replace_hyphen:
                    lyr_text = lyr_text.replace("-", "\u2212")
                rstr += lyr_text
        with Path(path).open("w", newline=self.options.newline) as f:
            f.write(rstr)
        logger.info(f"exported txt file: {Path(path).absolute()}")
        return self

    def export_docx(self, path: Path):
        import docx
        from docx.document import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Mm, Pt

        document: Document = docx.Document()
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(30)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(12.7)
        document.styles["Normal"].font.name = "Times New Roman"
        document.styles["Normal"].font.size = Pt(10)
        document.styles["Normal"].paragraph_format.space_after = Pt(0)
        paragraph = document.add_paragraph()
        for c in self.get():
            for lyr in self.layers:
                if isinstance(lyr, Layers.Break):
                    if lyr.break_type == "line":
                        paragraph = document.add_paragraph()
                        continue
                    break_key = {
                        "page": WD_BREAK.PAGE,
                        "column": WD_BREAK.COLUMN,
                        "section": WD_BREAK.PAGE,
                    }.get(lyr.break_type)
                    paragraph.add_run().add_break(break_key)
                    continue
                lyr_text = lyr.get_str(c)
                if lyr.replace_hyphen:
                    lyr_text = lyr_text.replace("-", "\u2212")
                run = paragraph.add_run(lyr_text)
                if lyr.bold:
                    run.bold = True
                if lyr.italic:
                    run.italic = True
                if lyr.underline:
                    run.underline = True
                if lyr.superscript:
                    run.font.superscript = True
                if lyr.subscript:
                    run.font.subscript = True
        path = Path(path).with_suffix(".docx")
        document.save(str(path))
        logger.info(f"exported docx file: {path.absolute()}")
        return self
